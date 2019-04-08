import os
import string
from pathlib import Path
import re

from docx import Document

from config import DATA_PATH, TABLES_TITLES
from textprocessing import table_header_clean, table_title_clean


"""
Получение из текста заголовков таблиц (заданы в config.py)
"""
def get_tables_titles(prgrs):
    tables_names = []
    for prgrp in prgrs:
        for name in TABLES_TITLES:
            if name in table_title_clean(prgrp.text):
                tables_names.append(name)
    
    return tables_names

"""
Получение названий колонок для таблицы.
Есть обработка случая двух колонок с одинаковыми 
названиями (до трех еще не добрался)
"""
def get_current_header(table):
    header = []
    for cell in table.rows[0].cells:
        cell_value = table_header_clean(cell.text)
        if cell_value in header:
            header.append(f'_{cell_value}')
        else:
            header.append(cell_value)
    return header

"""
Заполнение общего списка названий колонок
"""
def extract_header(table, header):
    current_header = get_current_header(table)
    for cell in current_header:
        if cell not in header:
            header.append(cell)
    print(current_header)
    return header

def save_table_header(file_path, tables_headers):
    print(file_path)
    doc = Document(file_path)

    tables_names = get_tables_titles(doc.paragraphs)

    if len(doc.tables) == 3 and tables_names == ['общий анализ крови', 'биохимический анализ', 'общий анализ мочи']:
        tables_headers['table1'] = extract_header(doc.tables[0], tables_headers['table1'])
        tables_headers['table2'] = extract_header(doc.tables[1], tables_headers['table2'])
        tables_headers['table3'] = extract_header(doc.tables[2], tables_headers['table3'])
    elif len(doc.tables) == 2 and tables_names == ['общий анализ крови', 'общий анализ мочи']:
        tables_headers['table1'] = extract_header(doc.tables[0], tables_headers['table1'])
        tables_headers['table3'] = extract_header(doc.tables[1], tables_headers['table3'])
    elif len(doc.tables) == 2 and tables_names == ['общий анализ крови', 'биохимический анализ']:
        tables_headers['table1'] = extract_header(doc.tables[0], tables_headers['table1'])
        tables_headers['table2'] = extract_header(doc.tables[1], tables_headers['table2'])
    else:
        print('Таблиц меньше 2')

    return tables_headers

def main():
    tables_headers = {
        'table1': [],
        'table2': [],
        'table3': []
    }

    for rootdir, dirs, files in os.walk(DATA_PATH):
        for filename in files:
            if filename.endswith('.docx'):
                file_path = Path(rootdir, filename)
                tables_headers = save_table_header(file_path, tables_headers)
    print(tables_headers['table1'])
    print(tables_headers['table2'])
    print(tables_headers['table3'])
    # i = 0
    # for table in TABLE_LIST:
    #     save_table_header(table, i)
    #     i += 1
    # sandbox()


if __name__ == "__main__":
    main()