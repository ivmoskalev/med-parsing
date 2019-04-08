from pathlib import Path
import os
import re
import datetime

from docx import Document
import pandas as pd

from textprocessing import (diseas_matcher, get_patient_id, date_processing, 
    get_patient_birth_date, get_patient_adress, approx_birth_date)
from tableprocessing import generate_empty_table_data, get_table_data
from tableheader import get_tables_titles
from config import DISEAS_DICT, DATA_PATH, TEST_PATH, TEXT_PATTERNS, TABLES_TITLES
from dbprocessing import desease_match


def get_file_objects(file_path):
    doc = Document(file_path)
    paragraphs = doc.paragraphs
    tables = doc.tables
    return paragraphs, tables

def paragraph_match(paragraphs, pattern):
    for prgrp in paragraphs:
        if pattern in prgrp.text:
            return(prgrp.text)
    return None

def extract_data_from_text(paragraphs):
    result = {}
    for key, value in TEXT_PATTERNS.items():
        mached_paragraph = paragraph_match(paragraphs, value)
        if mached_paragraph:
            if key in ['entrance_date', 'leave_date']:
                result_value = date_processing(mached_paragraph)
            elif key == 'id':
                result_value =  get_patient_id(mached_paragraph)
            elif key == 'birth_date':
                result_value = get_patient_birth_date(mached_paragraph)
            elif key == 'adress':
                result_value = get_patient_adress(mached_paragraph)
            # elif key == 'diagnosis':
            #     result_value = diseas_matcher(mached_paragraph)
            #     if all(v == 0 for v in result_value.values()):
            #         result_value = {
            #             'diagnos': 'Нет диагнозов',
            #             'text': mached_paragraph,
            #         }
            else:
                result_value = 'нет обработчика'
        else:
            result_value = 'нет совпадения в тексте'
        if type(result_value) == dict:
            result.update(result_value)
        else:
            result[key] = result_value
    if (type(result['birth_date']) is not datetime.date) and result['birth_date'] and type(result['entrance_date']) is datetime.date:
        result['birth_date'] = approx_birth_date(result['entrance_date'], result['birth_date'])

    return result

def extract_data_from_tables(tables, paragraphs, entrance_date):
    tables_data = generate_empty_table_data()
    tables_names = get_tables_titles(paragraphs)

    if type(entrance_date) is datetime.date:
        if len(tables) == 3 and tables_names == ['общий анализ крови', 'биохимический анализ', 'общий анализ мочи']:
            tables_data['table1'] = get_table_data(tables[0], tables_data['table1'], entrance_date)
            tables_data['table2'] = get_table_data(tables[1], tables_data['table2'], entrance_date)
            tables_data['table3'] = get_table_data(tables[2], tables_data['table3'], entrance_date)

        elif len(tables) == 2 and tables_names == ['общий анализ крови', 'общий анализ мочи']:
            tables_data['table1'] = get_table_data(tables[0], tables_data['table1'], entrance_date)
            tables_data['table2'] = get_table_data(tables[1], tables_data['table2'], entrance_date)
        elif len(tables) == 2 and tables_names == ['общий анализ крови', 'биохимический анализ']:
            tables_data['table1'] = get_table_data(tables[0], tables_data['table1'], entrance_date)
            tables_data['table3'] = get_table_data(tables[1], tables_data['table3'], entrance_date)

    """
    todo: работа с таблицами
    """
    # tables_names.append(len(tables))
    return tables_data

def extract_data(file_path):
    paragraphs, tables = get_file_objects(file_path)
    # print(file_path, len(tables))
    text_data = extract_data_from_text(paragraphs)
    table_data = extract_data_from_tables(tables, paragraphs, text_data['entrance_date'])
    # print(table_data)

    # for key, value in text_data.items():
    #     print('{}: {}'.format(key, value))
    # for key, value in table_data.items():
    #     print('{}: {}'.format(key, value))
    for key in table_data.keys():
        data = {f'{key}_{k}': v for k, v in table_data[key].items()} 
        text_data.update(data)
    
    return text_data    

def main():
    """
    Перебор всех файлов в каталоге
    Получение полного пути файла
    Инициация извлечения данных
    """
    # os.chdir(DATA_PATH)
    result = []
    # for filename in os.listdir(os.getcwd()):
    for rootdir, dirs, files in os.walk(DATA_PATH):
        for filename in files:
            if filename.endswith('.docx'):
                file_path = Path(rootdir, filename)
                # print(file_path)
                file_data = extract_data(file_path)
                result.append(file_data) 
    # print(result)
    # csv_path_data = Path(Path(DATA_PATH).parent, 'output','new_data.csv')
    # csv_path_disease = Path(Path(DATA_PATH).parent, 'output','new_disease.csv')
    csv_path_data = Path(Path(DATA_PATH).parent, 'output','test_data.csv')
    # csv_path_disease = Path(Path(DATA_PATH).parent, 'output','new_disease.csv')
    df = pd.DataFrame(data=result, index=None)
    disease_df = desease_match(df['id'])
    # print(df.head())
    # print(disease_df.head())
    disease_df = pd.merge(df, disease_df, how='left')
    disease_df.to_csv(csv_path_data)   

if __name__ == "__main__":
    main()
        
