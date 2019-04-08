import os
from pathlib import Path

import pandas as pd
import psycopg2
from docx import Document

from config import DATA_PATH, DB_CONFIG, TEXT_PATTERNS, DISEAS_DICT

def db_connect():
    return psycopg2.connect(**DB_CONFIG)

def diseas_parser(diseas):
    
    temp_list = []
    for operand in diseas[0]:
        if type(operand) == tuple:
            temp_list.append(diseas_parser(operand))
        else:
            temp_list.append(operand)
    separator = ' ' + diseas[1] + ' '
    result = '(' + separator.join(temp_list) + ')'
    return result

def get_query_string(diseas):
    
    query_string = "SELECT patient FROM diagnosis, \
        to_tsquery('russian', '{}') query \
        WHERE query @@ meta_diagnos \
        AND ts_rank_cd(meta_diagnos, query) >= 0.05;".format(diseas_parser(diseas))
    
    return query_string

def desease_match(base_df):
    
    with psycopg2.connect(**DB_CONFIG) as conn:
        cur = conn.cursor()
        for key, value in DISEAS_DICT.items():
            query_string = get_query_string(value)
            cur.execute(query_string)
            if cur.rowcount > 0:
                disease_df =  pd.DataFrame(
                        {'id': [row[0] for row in cur.fetchall()],
                        'd{}'.format(key): [1]*cur.rowcount,}
                    )
                base_df = pd.merge(base_df, disease_df, how='left')
            else:
                base_df['d{}'.format(key)] = 0.0
        cur.close()
        return base_df.fillna(0)

    

def get_file_objects(file_path):
    doc = Document(file_path)
    paragraphs = doc.paragraphs
    tables = doc.tables
    return paragraphs

def paragraph_match(paragraphs, pattern):
    for prgrp in paragraphs:
        if pattern in prgrp.text:
            return(prgrp.text)
    return None

def get_text_without_key(paragraphs, pattern):
    result = ''
    for prgrp in paragraphs:
        if pattern not in prgrp.text:
            if result == '':
                result += prgrp.text
            else:
                result += f'\n{prgrp.text}'
    return result

def extract_data_from_text(paragraphs):
    result = {}
    for key in ['id', 'diagnosis']:
        mached_paragraph = paragraph_match(paragraphs, TEXT_PATTERNS[key])
        if mached_paragraph:
            if key == 'id':
                try:
                    result_value =  mached_paragraph.split(':')[1].strip()
                except:
                    print('Не удалось обработать id пациета')
                    return None
            elif key == 'diagnosis':
                result_value = mached_paragraph
                other_text = get_text_without_key(paragraphs, TEXT_PATTERNS[key])
        else:
            result_value = 'Нет совпадения в тексте'
            return None
        result[key] = result_value
    result['other_text'] = other_text
    return result

def save_data(file_path, db_cur):
    paragraphs = get_file_objects(file_path)
    print(file_path)
    text_data = extract_data_from_text(paragraphs)
    if text_data:
        db_cur.execute("INSERT INTO diagnosis (patient, diagnos, other_text) VALUES (%s, %s, %s)", (text_data['id'], text_data['diagnosis'], text_data['other_text']))
    else:
        print('WARNNING!!! {}'.format(file_path))
        return    

def update_data(file_path, db_cur):
    paragraphs = get_file_objects(file_path)
    print(file_path)
    other_text = get_text_without_key(paragraphs, TEXT_PATTERNS['diagnosis'])
    patient_id = paragraph_match(paragraphs, TEXT_PATTERNS['id']).split(':')[1].strip()
    if other_text != '':
        # print(other_text)
        db_cur.execute("UPDATE diagnosis SET other_text = (%s) WHERE patient = (%s)", (other_text, patient_id))
    else:
        print('WARNNING!!! {}',format(file_path))
        return


def file_iterator(path, db_cur):
    for rootdir, dirs, files in os.walk(path):
        for filename in files:
            if filename.endswith('.docx'):
                file_path = Path(rootdir, filename)
                save_data(file_path, db_cur)
                # update_data(file_path, db_cur)


def main():
    """
    Перебор всех файлов в каталоге
    Получение полного пути файла
    Инициация извлечения данных
    """
    # os.chdir(DATA_PATH)
    # for filename in os.listdir(os.getcwd()):
    with psycopg2.connect(**DB_CONFIG) as conn:
        cur = conn.cursor()
        desease_match(cur)
        # file_iterator(DATA_PATH, cur)
        conn.commit()
        cur.close()

if __name__ == "__main__":
    main()
        